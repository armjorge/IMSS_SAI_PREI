import os
from config import ConfigManager
from sql_connexion_updating import SQL_CONNEXION_UPDATING
from sqlalchemy import create_engine, text
import pandas as pd
from datetime import datetime
from typing import Optional

try:
    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_pdf import PdfPages
    from matplotlib.ticker import FuncFormatter  # Movido aquí para evitar reimports
    _HAS_MPL = True
except Exception:
    _HAS_MPL = False

try:
    from docx import Document
    from docx.shared import Inches
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False


class DataWarehouse:
    def __init__(self, data_access, working_folder=None):
        self.data_access = data_access
        self.working_folder = working_folder or os.getcwd()
        
    def split_df_by_date(self, dataframe, cutoff_date, ciclo):
        estado_col = None
        for c in dataframe.columns:
            if 'estado' in c.lower():
                estado_col = c
                break
        if estado_col is None:
            raise ValueError("No se encontró la columna de estado")
        print(f"Usando columna de estado: {estado_col}")

        df_tycsa = dataframe[dataframe['fechaaltatrunc'] < cutoff_date]
        df_cpi = dataframe[dataframe['fechaaltatrunc'] >= cutoff_date]
         
        grouped_raw = dataframe.groupby(estado_col)['importe'].sum()
        grouped_dftycsa = df_tycsa.groupby(estado_col)['importe'].sum()
        grouped_df_cpi = df_cpi.groupby(estado_col)['importe'].sum()

        print(f'RAW AGRUPADO {ciclo}\n', grouped_raw.head())
        print(f'DF TYCSA AGRUPADO {ciclo}\n', grouped_dftycsa.head())
        print(f'DF CPI AGRUPADO {ciclo}\n', grouped_df_cpi.head())
        
        # Return the grouped Series instead of DataFrames
        return grouped_df_cpi, grouped_dftycsa, grouped_raw
    
    def generate_altas_historico_report(self, df_altas_historico: pd.DataFrame,
                                        report_folder: Optional[str] = None) -> Optional[str]:
        print("Inicio de generate_altas_historico_report")  # Print de depuración
        """
        Genera un DOCX con secciones para PTYCSA y CPI:
        - Filtra df_altas_historico por 'fechaaltatrunc' (< 2025-06-30 para PTYCSA, >= para CPI).
        - Para cada subconjunto: gráfico de barras comparativo, tabla resumen y gráfico de tendencias.
        - Selección interactiva de fechas aplicada al DataFrame completo.
        - Guarda en report_folder/consulta {YYYY} {MM} {DD}.docx
        """
        if df_altas_historico is None or df_altas_historico.empty:
            print("Sin datos para reporte")
            return None

        df = df_altas_historico.copy()
        df.info()

        # Verificar y convertir 'fechaaltatrunc' a datetime, luego a date (para coincidir con consulta SQL)
        if 'fechaaltatrunc' not in df.columns:
            raise ValueError("Se requiere 'fechaaltatrunc' para filtrar PTYCSA y CPI")
        df['fechaaltatrunc'] = pd.to_datetime(df['fechaaltatrunc'], errors='coerce').dt.date

        # Tipos para columnas comunes
        if 'file_date' not in df.columns:
            raise ValueError("Se requiere 'file_date'")
        df['file_date'] = pd.to_datetime(df['file_date'], errors='coerce')  # Keep full datetime including hour
        if 'importe' not in df.columns:
            raise ValueError("Se requiere 'importe'")
        df['importe'] = pd.to_numeric(df['importe'], errors='coerce').fillna(0)


        # Detectar columna de estado (usando df completo, asumiendo consistencia)
        """
        def _norm(s: str) -> str:
            return ''.join(ch for ch in s.lower() if ch.isalnum() or ch == '_')
        estado_col = None
        for c in df.columns:
            if _norm(c) == 'estado_cr':
                estado_col = c
                break
        if estado_col is None and 'estado_c.r.' in df.columns:
            estado_col = 'estado_c.r.'
        if estado_col is None:
            raise ValueError("No se encontró la columna de estado (esperado 'estado_c.r.')")
        """
        # Fechas disponibles (del DataFrame completo)
        dates = sorted(df['file_date'].dropna().unique())
        print(f"file_date.nunique = {len(dates)}")
        if not dates:
            print("No hay fechas válidas")
            return None

        # Interactive selection loop (una vez, para ambos subconjuntos)
        while True:
            print("\nFechas disponibles:")
            for i, date in enumerate(dates):
                print(f"{i}: {date}")
            try:
                current_idx = int(input("Elige índice para reporte_a_comparar (current): "))
                prev_idx = int(input("Elige índice para reporte_previo (previous): "))
                if 0 <= current_idx < len(dates) and 0 <= prev_idx < len(dates):
                    current_date = dates[current_idx]
                    prev_date = dates[prev_idx]
                    print(f"Seleccionado - Current: {current_date}, Previous: {prev_date}")
                    confirm = input("Confirmar? (y/n): ").lower()
                    if confirm == 'y':
                        break
                else:
                    print("Índices inválidos.")
            except ValueError:
                print("Entrada inválida. Usa números enteros.")

        prev_label = f"Seleccionado: {prev_date}"

        # Filtrar DataFrames por cortes (usando .date() para coincidir con consulta SQL)


        cutoff_date = pd.to_datetime('2025-06-30').date()
        df_previous = df[df['file_date'] == prev_date]
        df_current = df[df['file_date'] == current_date]

        # Split for previous
        grouped_cpi_prev, grouped_tycsa_prev, grouped_raw_prev = self.split_df_by_date(df_previous, cutoff_date, prev_date)
        # Split for current
        grouped_cpi_curr, grouped_tycsa_curr, grouped_raw_curr = self.split_df_by_date(df_current, cutoff_date, current_date)
        # Merge into summary DataFrames with dates as columns
        summary_raw = pd.concat([grouped_raw_prev.rename(prev_date), grouped_raw_curr.rename(current_date)], axis=1).fillna(0)
        summary_tycsa = pd.concat([grouped_tycsa_prev.rename(prev_date), grouped_tycsa_curr.rename(current_date)], axis=1).fillna(0)
        summary_cpi = pd.concat([grouped_cpi_prev.rename(prev_date), grouped_cpi_curr.rename(current_date)], axis=1).fillna(0)


        # Crear y imprimir tablas resumen
        print("\n=== Summary RAW (previous y current combinados) ===")
        # Calcular delta y delta_pct
        summary_raw['delta'] = summary_raw[current_date] - summary_raw[prev_date]
        summary_raw['delta_pct'] = summary_raw.apply(lambda r: (r['delta']/r[prev_date]*100.0) if r[prev_date] else None, axis=1)
        # Add total_period column
        summary_raw['total_period'] = summary_raw[prev_date] + summary_raw[current_date]
        # Agregar fila de Total
        summary_raw.loc['Total', prev_date] = summary_raw[prev_date].sum()
        summary_raw.loc['Total', current_date] = summary_raw[current_date].sum()
        summary_raw.loc['Total', 'delta'] = summary_raw['delta'].sum()
        if summary_raw.loc['Total', prev_date] != 0:
            summary_raw.loc['Total', 'delta_pct'] = (summary_raw.loc['Total', 'delta'] / summary_raw.loc['Total', prev_date]) * 100
        else:
            summary_raw.loc['Total', 'delta_pct'] = None
        summary_raw.loc['Total', 'total_period'] = summary_raw['total_period'].sum()
        print(summary_raw)

        print("\n=== Summary PTYCSA (previous y current combinados) ===")
        # Calcular delta y delta_pct
        summary_tycsa['delta'] = summary_tycsa[current_date] - summary_tycsa[prev_date]
        summary_tycsa['delta_pct'] = summary_tycsa.apply(lambda r: (r['delta']/r[prev_date]*100.0) if r[prev_date] else None, axis=1)
        # Add total_period column
        summary_tycsa['total_period'] = summary_tycsa[prev_date] + summary_tycsa[current_date]
        # Agregar fila de Total
        summary_tycsa.loc['Total', prev_date] = summary_tycsa[prev_date].sum()
        summary_tycsa.loc['Total', current_date] = summary_tycsa[current_date].sum()
        summary_tycsa.loc['Total', 'delta'] = summary_tycsa['delta'].sum()
        if summary_tycsa.loc['Total', prev_date] != 0:
            summary_tycsa.loc['Total', 'delta_pct'] = (summary_tycsa.loc['Total', 'delta'] / summary_tycsa.loc['Total', prev_date]) * 100
        else:
            summary_tycsa.loc['Total', 'delta_pct'] = None
        summary_tycsa.loc['Total', 'total_period'] = summary_tycsa['total_period'].sum()
        print(summary_tycsa)

        print("\n=== Summary CPI (previous y current combinados) ===")
        # Calcular delta y delta_pct
        summary_cpi['delta'] = summary_cpi[current_date] - summary_cpi[prev_date]
        summary_cpi['delta_pct'] = summary_cpi.apply(lambda r: (r['delta']/r[prev_date]*100.0) if r[prev_date] else None, axis=1)
        # Add total_period column
        summary_cpi['total_period'] = summary_cpi[prev_date] + summary_cpi[current_date]
        # Agregar fila de Total
        summary_cpi.loc['Total', prev_date] = summary_cpi[prev_date].sum()
        summary_cpi.loc['Total', current_date] = summary_cpi[current_date].sum()
        summary_cpi.loc['Total', 'delta'] = summary_cpi['delta'].sum()
        if summary_cpi.loc['Total', prev_date] != 0:
            summary_cpi.loc['Total', 'delta_pct'] = (summary_cpi.loc['Total', 'delta'] / summary_cpi.loc['Total', prev_date]) * 100
        else:
            summary_cpi.loc['Total', 'delta_pct'] = None
        summary_cpi.loc['Total', 'total_period'] = summary_cpi['total_period'].sum()
        print(summary_cpi)
        # Rename columns for consistency in generate_summary_section
        summary_tycsa = summary_tycsa.rename(columns={prev_date: 'previous', current_date: 'current'})
        summary_cpi = summary_cpi.rename(columns={prev_date: 'previous', current_date: 'current'})

        # add 'Total' and the entire period column for summary_cpi, summary_tycsa, summary_raw
        # Función auxiliar para generar sección por summary
        def generate_summary_section(doc, summary, subset_name, current_date, prev_date, prev_label, out_dir):
            print(f"Generando sección {subset_name}")  # Print de depuración
            if summary.empty:
                print(f"Summary {subset_name} vacío, omitiendo sección.")
                return

            # Calcular delta y delta_pct (si no se ha hecho ya)
            if 'delta' not in summary.columns:
                summary['delta'] = summary['current'] - summary['previous']
                summary['delta_pct'] = summary.apply(lambda r: (r['delta']/r['previous']*100.0) if r['previous'] else None, axis=1)

            # Separar fila de Total para colocarla al final
            total_row = summary.loc[['Total']] if 'Total' in summary.index else None
            summary = summary.drop('Total', errors='ignore')

            # Ordenar el resto por 'current' descendente
            summary = summary.sort_values('current', ascending=False)

            # Agregar Total al final
            if total_row is not None:
                summary = pd.concat([summary, total_row])

            # Formatting
            def format_currency(x):
                return f"${x:,.2f}" if pd.notnull(x) else ""
            summary_formatted = summary.copy()
            for col in ['current', 'previous', 'delta']:
                summary_formatted[col] = summary_formatted[col].apply(format_currency)
            summary_formatted['delta_pct'] = summary_formatted['delta_pct'].apply(lambda x: f"{x:.2f}%" if pd.notnull(x) else "")

            # Subtítulo
            doc.add_heading(f"Sección {subset_name}", level=1)

            # Página: barras comparativas
            fig, ax = plt.subplots(figsize=(11.69, 8.27))
            idx = summary.index.tolist()
            x = range(len(idx))
            ax.bar([i - 0.2 for i in x], summary['previous'].values, width=0.4, label='Previous')
            ax.bar([i + 0.2 for i in x], summary['current'].values, width=0.4, label='Current')
            ax.set_xticks(list(x))
            ax.set_xticklabels(idx, rotation=45, ha='right')
            ax.set_ylabel('Importe')
            ax.set_title(f"Comparativo {subset_name} - current: {current_date} | prev: {prev_label}")
            ax.legend()
            ax.grid(axis='y', linestyle='--', alpha=0.3)
            ax.yaxis.set_major_formatter(FuncFormatter(lambda x, _: f"${x:,.0f}"))
            fig.tight_layout()
            bar_img_path = os.path.join(out_dir, f'temp_bar_{subset_name}.png')
            fig.savefig(bar_img_path)
            plt.close(fig)
            doc.add_picture(bar_img_path, width=Inches(6))

            # Página: tabla
            doc.add_heading(f'Resumen por estado {subset_name} (top)', level=2)
            table = doc.add_table(rows=1, cols=len(summary_formatted.columns) + 1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Estado'
            hdr_cells[1].text = str(prev_date)  # Just the date
            hdr_cells[2].text = str(current_date)  # Just the date
            hdr_cells[3].text = 'Delta'
            hdr_cells[4].text = 'Delta %'
            for i, row in enumerate(summary_formatted.head(30).itertuples(index=True)):
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.Index)
                row_cells[1].text = str(row.previous)
                row_cells[2].text = str(row.current)
                row_cells[3].text = str(row.delta)
                row_cells[4].text = str(row.delta_pct)

            # Limpiar imagen
            try:
                os.remove(bar_img_path)
                print(f"Imagen {bar_img_path} eliminada.")
            except Exception as e:
                print(f"Error eliminando imagen {bar_img_path}: {e}")

            print(f"Sección {subset_name} generada exitosamente.")  # Print de depuración

        # Preparar carpeta y archivo
        out_dir = report_folder or os.path.join(self.working_folder, 'Reportes BI')
        os.makedirs(out_dir, exist_ok=True)
        today = datetime.now()
        out_docx = os.path.join(out_dir, f"consulta {today.year} {today.month:02d} {today.day:02d}.docx")

        title = f"Avance de Contrarecibos en el sistema PREI - PTYCSA y CPI - current: {current_date} | prev: {prev_label}"

        print(f"_HAS_DOCX: {_HAS_DOCX}, _HAS_MPL: {_HAS_MPL}")  # Print de depuración

        if not _HAS_DOCX:
            print("python-docx no disponible. Generando CSVs en su lugar.")
            summary_tycsa.to_csv(os.path.join(out_dir, f"consulta_{today.year}{today.month:02d}{today.day:02d}_PTYCSA_summary.csv"))
            summary_cpi.to_csv(os.path.join(out_dir, f"consulta_{today.year}{today.month:02d}{today.day:02d}_CPI_summary.csv"))
            return None

        if not _HAS_MPL:
            print("Matplotlib no disponible. No se pueden generar gráficos.")
            return None

        try:
            print("Creando documento DOCX...")  # Print de depuración
            doc = Document()
            doc.add_heading(title, 0)

            # Agregar secciones para PTYCSA y CPI
            generate_summary_section(doc, summary_tycsa, "PTYCSA", current_date, prev_date, prev_label, out_dir)
            generate_summary_section(doc, summary_cpi, "CPI", current_date, prev_date, prev_label, out_dir)

            print(f"Guardando DOCX en: {out_docx}")  # Print de depuración
            doc.save(out_docx)
            print(f"Reporte generado: {out_docx}")
            return out_docx
        except Exception as e:
            print(f"Error generando reporte DOCX: {e}")
            return None 

        

    def Business_Intelligence(self):
        source_schema = "eseotres_warehouse"
        #user_input = input('Elige la base del análisis, 1) cortes jupyter lab (ciclos fiscales completos), 2) cortes mini imss (sólo 6 junio): ')
        source_table =  "altas_historicas"
        print(f"📦 Fuente: {source_schema}.{source_table}")
        print("Conectando a la base de datos SOURCE...")
        #print(self.data_access)
        # Leer desde diccionario (evita AttributeError si no es objeto)
        sql_url = self.data_access.get('sql_url') if isinstance(self.data_access, dict) else None
        target_url = self.data_access.get('sql_target') if isinstance(self.data_access, dict) else None

        if not sql_url:
            print("❌ 'sql_url' no encontrado en data_access")
            return

        # Probar conexión
        try:
            engine = create_engine(sql_url)
            with engine.connect() as conn:
                ok = conn.execute(text('SELECT 1')).scalar()
                ver = conn.execute(text('SELECT version()')).scalar()
                # Cargar tabla origen como df_source
                query = f'SELECT * FROM "{source_schema}"."{source_table}"'
                df_source = pd.read_sql_query(text(query), conn)
                self.df_source = df_source
                try:
                    print(f"📊 df_source cargado: {df_source.shape[0]} filas, {df_source.shape[1]} columnas")
                    print("Llamando a generate_altas_historico_report")  # Print de depuración
                    self.generate_altas_historico_report(df_source)  # Updated call, removed previous_mode
                except Exception as e:
                    print(f"Error en generate_altas_historico_report: {e}")  # Imprimir error real
                print(f"✅ Conexión OK (SELECT 1 => {ok})")
                print(f"🗄️ Versión servidor: {ver}")
        except Exception as e:
            print(f"❌ Error al conectar/probar la BD: {e}")
            return
        finally:
            try:
                engine.dispose()
            except Exception:
                pass

        print("Inteligencia de negocios")




if __name__ == "__main__":
    folder_root = os.getcwd()
    working_folder = os.path.join(folder_root, "Implementación")
    config_manager = ConfigManager(working_folder)
    data_access = config_manager.yaml_creation(working_folder)
    data_warehouse = DataWarehouse(data_access, working_folder)
    data_warehouse.Business_Intelligence()
